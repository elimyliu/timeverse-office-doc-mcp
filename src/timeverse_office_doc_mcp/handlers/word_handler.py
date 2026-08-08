"""Word 文档处理器 - 18 个工具。

对应方案 5.1 Word 工具集。
使用 python-docx 实现，支持 Session 内存编辑模式。
"""

from __future__ import annotations

import logging
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from ..common.error_handler import ToolError
from ..common.file_lock import file_lock_mgr
from ..common.path_guard import path_guard
from ..common.session import session_manager
from ..common.template_mgr import template_manager
from ..common.template_utils import fill_word_variables
from ..common.validator import InputValidator

logger = logging.getLogger("timeverse_office_doc_mcp.word")


# ==================== 辅助函数 ====================


def _get_document(filename: str, session_id: str | None = None) -> Document:
    """获取 Document 对象：Session 模式从内存取，否则从磁盘打开。"""
    if session_id:
        return session_manager.get_document(session_id, "word")
    validated = path_guard.validate_path(filename, "read")
    return Document(validated)


def _save_document(doc: Document, filename: str, session_id: str | None = None) -> None:
    """保存文档：Session 模式仅标记修改，否则写入磁盘。"""
    if session_id:
        session_manager.mark_modified(session_id)
    else:
        validated = path_guard.validate_path(filename, "write")
        file_lock_mgr.acquire(validated)
        try:
            doc.save(validated)
        finally:
            file_lock_mgr.release(validated)


# ==================== 5.1.1 文档管理（5 个） ====================


def word_create_document(
    filename: str,
    title: str = "",
    author: str = "",
    template: str | None = None,
    variables: dict[str, Any] | None = None,
    session_id: str | None = None,
) -> dict[str, Any]:
    """创建新文档（支持模板）。"""
    InputValidator.validate_filename(filename)
    validated = path_guard.validate_path(filename, "write")

    if template:
        # 从模板创建
        fmt, tpl_path = template_manager.resolve_template_path(template)
        if fmt != "word":
            raise ToolError(f"模板 '{template}' 是 {fmt} 格式，不是 word")
        doc = Document(tpl_path)
        replaced = fill_word_variables(doc, variables) if variables else 0
    else:
        doc = Document()
        replaced = 0

    # 设置核心属性
    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    if session_id:
        session_manager.open_session(validated, "word", doc)
    else:
        file_lock_mgr.acquire(validated)
        try:
            doc.save(validated)
        finally:
            file_lock_mgr.release(validated)

    return {
        "filename": validated,
        "title": title,
        "author": author,
        "template": template,
        "variables_replaced": replaced,
        "session_id": session_id,
    }


def word_get_info(
    filename: str,
    detailed: bool = False,
    session_id: str | None = None,
) -> dict[str, Any]:
    """获取文档元信息（detailed=True 时附带结构分析）。"""
    doc = _get_document(filename, session_id)
    props = doc.core_properties
    result: dict[str, Any] = {
        "filename": filename,
        "title": props.title or "",
        "author": props.author or "",
        "subject": props.subject or "",
        "created": props.created.isoformat() if props.created else None,
        "modified": props.modified.isoformat() if props.modified else None,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
    }

    if detailed:
        # 标题层级统计
        heading_levels: dict[int, int] = {}
        style_distribution: dict[str, int] = {}
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"
            style_distribution[style_name] = style_distribution.get(style_name, 0) + 1
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                    heading_levels[level] = heading_levels.get(level, 0) + 1
                except ValueError:
                    heading_levels[1] = heading_levels.get(1, 0) + 1

        # 表格统计
        table_details: list[dict[str, int]] = []
        for table in doc.tables:
            table_details.append({"rows": len(table.rows), "cols": len(table.columns)})

        # 图片统计
        image_count = 0
        for _shape in doc.inline_shapes:
            image_count += 1

        total_chars = sum(len(p.text) for p in doc.paragraphs)

        result.update(
            {
                "heading_levels": heading_levels,
                "style_distribution": style_distribution,
                "table_details": table_details,
                "image_count": image_count,
                "total_chars": total_chars,
            }
        )

    return result


def word_extract_text(
    filename: str,
    extract_type: str = "text",
    include_tables: bool = True,
    session_id: str | None = None,
) -> dict[str, Any]:
    """提取全文文本或大纲结构。"""
    InputValidator.validate_choice(extract_type, ["text", "outline", "all"], "extract_type")
    doc = _get_document(filename, session_id)
    result: dict[str, Any] = {"filename": filename}

    if extract_type in ("text", "all"):
        parts: list[str] = [p.text for p in doc.paragraphs]
        table_texts: list[list[list[str]]] = []
        if include_tables:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                table_texts.append(rows)
        result["text"] = "\n".join(parts)
        result["paragraph_count"] = len(doc.paragraphs)
        result["tables"] = table_texts if include_tables else []

    if extract_type in ("outline", "all"):
        outline: list[dict[str, Any]] = []
        for idx, para in enumerate(doc.paragraphs):
            if para.style and para.style.name.startswith("Heading"):
                level = 0
                style_name = para.style.name
                if style_name == "Title":
                    level = 0
                else:
                    try:
                        level = int(style_name.split()[-1])
                    except ValueError:
                        level = 1
                outline.append(
                    {
                        "paragraph_idx": idx,
                        "level": level,
                        "text": para.text,
                        "style": style_name,
                    }
                )
        result["outline"] = outline

    return result


def word_list_documents(directory: str) -> dict[str, Any]:
    """列出目录内文档。"""
    validated_dir = path_guard.validate_directory(directory)
    docx_files = sorted(Path(validated_dir).glob("*.docx"))
    return {
        "directory": validated_dir,
        "documents": [
            {"filename": str(f), "name": f.name, "size": f.stat().st_size} for f in docx_files
        ],
        "count": len(docx_files),
    }


def word_copy_document(source: str, destination: str) -> dict[str, Any]:
    """复制文档。"""
    src = path_guard.validate_path(source, "read")
    dst = path_guard.validate_path(destination, "write")
    file_lock_mgr.acquire(dst)
    try:
        shutil.copy2(src, dst)
    finally:
        file_lock_mgr.release(dst)
    return {"source": src, "destination": dst}


# ==================== 5.1.2 内容编辑（7 个） ====================


def word_add_heading(
    filename: str,
    text: str,
    level: int = 1,
    align: str | None = None,
    font_size: int | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
    session_id: str | None = None,
) -> dict[str, Any]:
    """添加标题（支持对齐、字号覆盖与段前/段后间距）。"""
    InputValidator.validate_text_length(text)
    if level < 0 or level > 9:
        raise ToolError(f"标题级别必须在 0-9 之间，得到: {level}")
    doc = _get_document(filename, session_id)
    heading = doc.add_heading(text, level=level)
    if font_size:
        for run in heading.runs:
            run.font.size = Pt(font_size)
    _apply_paragraph_format(heading, align, space_before, space_after)
    _save_document(doc, filename, session_id)
    return {
        "filename": filename,
        "text": text,
        "level": level,
        "style": heading.style.name,
        "align": align,
        "font_size": font_size,
        "space_before": space_before,
        "space_after": space_after,
    }


def word_add_paragraph(
    filename: str,
    text: str = "",
    style: str | None = None,
    font_size: int | None = None,
    bold: bool = False,
    align: str | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
    page_break: bool = False,
    session_id: str | None = None,
) -> dict[str, Any]:
    """添加段落（可选分页符；支持对齐与段前/段后间距）。"""
    if not text and not page_break:
        raise ToolError("text 和 page_break 至少需要一个")
    if text:
        InputValidator.validate_text_length(text)
    doc = _get_document(filename, session_id)
    if page_break:
        doc.add_page_break()
    if text:
        para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        run = para.add_run(text)
        if font_size:
            run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
        _apply_paragraph_format(para, align, space_before, space_after)
    _save_document(doc, filename, session_id)
    return {
        "filename": filename,
        "text": text,
        "style": style,
        "align": align,
        "space_before": space_before,
        "space_after": space_after,
        "page_break": page_break,
        "paragraph_idx": len(doc.paragraphs) - 1,
    }


def word_add_cover(
    filename: str,
    title: str,
    subtitle: str = "",
    author: str = "",
    date: str = "",
    org: str = "",
    session_id: str | None = None,
) -> dict[str, Any]:
    """添加居中版式封面（标题/副标题/作者/日期/组织），完成后自动分页。"""
    if not title.strip():
        raise ToolError("title 不能为空")
    InputValidator.validate_text_length(title)
    for label, value in [("subtitle", subtitle), ("author", author), ("date", date), ("org", org)]:
        if value:
            InputValidator.validate_text_length(value)
    doc = _get_document(filename, session_id)

    def _cover_para(
        text: str,
        size: int,
        bold: bool = False,
        space_before: float = 0,
        space_after: float = 0,
    ) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        _apply_paragraph_format(para, align="center", space_before=space_before, space_after=space_after)

    # 顶部留白（将标题推向页面上 1/3 位置）
    _cover_para(" ", 12)
    _cover_para(" ", 12)
    # 标题（用普通段落而非 Heading，避免进入目录/大纲）
    _cover_para(title, 26, bold=True, space_before=120)
    # 副标题
    if subtitle:
        _cover_para(subtitle, 16, space_before=24)
    # 中部留白
    for _ in range(3):
        _cover_para(" ", 12)
    # 底部信息
    for text in (date, author, org):
        if text:
            _cover_para(text, 12, space_after=6)
    # 封面结束，自动分页
    doc.add_page_break()

    _save_document(doc, filename, session_id)
    return {
        "filename": filename,
        "title": title,
        "subtitle": subtitle,
        "author": author,
        "date": date,
        "org": org,
        "page_break_added": True,
    }


def word_add_table(
    filename: str,
    rows: int,
    cols: int,
    data: list[list[str]] | None = None,
    has_header: bool = True,
    session_id: str | None = None,
) -> dict[str, Any]:
    """添加表格。"""
    InputValidator.validate_positive_int(rows, "rows")
    InputValidator.validate_positive_int(cols, "cols")
    if data:
        InputValidator.validate_table_data(data)
    doc = _get_document(filename, session_id)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    if data:
        for i, row_data in enumerate(data):
            if i >= rows:
                break
            for j, cell_text in enumerate(row_data):
                if j >= cols:
                    break
                table.cell(i, j).text = str(cell_text)

    if has_header and data and rows > 0:
        for j in range(min(cols, len(data[0]))):
            cell = table.cell(0, j)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    _save_document(doc, filename, session_id)
    return {
        "filename": filename,
        "rows": rows,
        "cols": cols,
        "table_idx": len(doc.tables) - 1,
    }


def word_add_image(
    filename: str,
    image_path: str,
    width: float | None = None,
    session_id: str | None = None,
) -> dict[str, Any]:
    """插入图片。"""
    img = path_guard.validate_path(image_path, "read")
    if not Path(img).exists():
        raise ToolError(f"图片文件不存在: {image_path}")
    doc = _get_document(filename, session_id)
    kwargs: dict[str, Any] = {}
    if width:
        kwargs["width"] = Inches(width)
    doc.add_picture(img, **kwargs)
    _save_document(doc, filename, session_id)
    return {"filename": filename, "image_path": img, "width": width}


def word_add_list(
    filename: str,
    items: list[str],
    list_style: str = "List Bullet",
    session_id: str | None = None,
) -> dict[str, Any]:
    """添加列表。"""
    if not items:
        raise ToolError("列表项不能为空")
    for item in items:
        InputValidator.validate_text_length(item)
    doc = _get_document(filename, session_id)
    for item in items:
        doc.add_paragraph(item, style=list_style)
    _save_document(doc, filename, session_id)
    return {"filename": filename, "items_count": len(items), "list_style": list_style}


def word_set_header_footer(
    filename: str,
    header_text: str = "",
    footer_text: str = "",
    include_page_num: bool = False,
    session_id: str | None = None,
) -> dict[str, Any]:
    """设置页眉页脚。"""
    doc = _get_document(filename, session_id)
    for section in doc.sections:
        if header_text:
            header = section.header
            header.is_linked_to_previous = False
            if header.paragraphs:
                header.paragraphs[0].text = header_text
            else:
                header.add_paragraph(header_text)
        if footer_text or include_page_num:
            footer = section.footer
            footer.is_linked_to_previous = False
            if footer_text:
                if footer.paragraphs:
                    footer.paragraphs[0].text = footer_text
                else:
                    footer.add_paragraph(footer_text)
            if include_page_num:
                para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                run = para.add_run()
                fldChar1 = _make_element("w:fldChar", {"w:fldCharType": "begin"})
                instrText = _make_element("w:instrText", {}, "PAGE")
                fldChar2 = _make_element("w:fldChar", {"w:fldCharType": "end"})
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
    _save_document(doc, filename, session_id)
    return {
        "filename": filename,
        "header": header_text,
        "footer": footer_text,
        "page_num": include_page_num,
    }


def _make_element(tag: str, attrs: dict[str, str], text: str = "") -> Any:
    """创建 XML 元素（用于页码域代码）。"""
    elem = etree.SubElement(etree.Element("dummy"), qn(tag))
    for k, v in attrs.items():
        elem.set(qn(k), v)
    if text:
        elem.text = text
    return elem


_ALIGN_OPTIONS: dict[str, WD_ALIGN_PARAGRAPH] = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _apply_paragraph_format(
    para: Any,
    align: str | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
) -> None:
    """应用段落对齐与段前/段后间距（用于封面等排版场景）。"""
    if align:
        key = align.strip().lower()
        if key not in _ALIGN_OPTIONS:
            raise ToolError(f"对齐方式无效: {align}，可选: {', '.join(_ALIGN_OPTIONS)}")
        para.alignment = _ALIGN_OPTIONS[key]
    pf = para.paragraph_format
    if space_before is not None:
        if space_before < 0:
            raise ToolError("space_before 不能为负数")
        pf.space_before = Pt(space_before)
    if space_after is not None:
        if space_after < 0:
            raise ToolError("space_after 不能为负数")
        pf.space_after = Pt(space_after)


def _enable_update_fields_on_open(doc: Document) -> None:
    """写入 <w:updateFields/>, 使文档打开时自动更新所有域（含 TOC 目录）。

    不写此项时 Word 打开文档默认不执行域计算，目录只会显示占位文本，
    需用户手动右键「更新域」或按 F9 才会生成。
    按 OOXML 规范顺序插入（位于 compat 等元素之前），避免产生无效文档。
    """
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return
    update_fields = settings.makeelement(qn("w:updateFields"), {qn("w:val"): "true"})
    settings.insert_element_before(
        update_fields,
        "w:hdrShapeDefaults",
        "w:footnotePr",
        "w:endnotePr",
        "w:compat",
        "w:docVars",
        "w:rsids",
        "m:mathPr",
        "w:attachedSchema",
        "w:themeFontLang",
        "w:clrSchemeMapping",
    )


def word_generate_toc(
    filename: str,
    max_level: int = 3,
    styles: str | None = None,
    session_id: str | None = None,
) -> dict[str, Any]:
    """生成目录（Table of Contents）。"""
    doc = _get_document(filename, session_id)
    para = doc.add_paragraph()
    run = para.add_run()
    # 插入 TOC 域代码
    fldChar_begin = etree.SubElement(run._r, qn("w:fldChar"))
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = etree.SubElement(run._r, qn("w:instrText"))
    instrText.set(qn("xml:space"), "preserve")
    toc_opts = f' TOC \\o "1-{max_level}"'
    if styles:
        toc_opts += f' \\t "{styles}"'
    instrText.text = toc_opts

    fldChar_sep = etree.SubElement(run._r, qn("w:fldChar"))
    fldChar_sep.set(qn("w:fldCharType"), "separate")

    # 占位文本
    fldText = etree.SubElement(run._r, qn("w:t"))
    fldText.text = "右键点击此处选择「更新域」以生成目录"

    fldChar_end = etree.SubElement(run._r, qn("w:fldChar"))
    fldChar_end.set(qn("w:fldCharType"), "end")

    # 打开文档时自动更新所有域（含 TOC），避免用户手动按 F9
    _enable_update_fields_on_open(doc)

    _save_document(doc, filename, session_id)
    return {"filename": filename, "max_level": max_level, "note": "打开文档后目录将自动生成"}


# ==================== 5.1.3 格式化与操作（5 个） ====================


def word_format_text(
    filename: str,
    paragraph_idx: int,
    start: int,
    end: int,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    font: str | None = None,
    session_id: str | None = None,
) -> dict[str, Any]:
    """格式化文本片段。"""
    doc = _get_document(filename, session_id)
    if paragraph_idx < 0 or paragraph_idx >= len(doc.paragraphs):
        raise ToolError(f"段落索引超出范围: {paragraph_idx}（共 {len(doc.paragraphs)} 段）")

    para = doc.paragraphs[paragraph_idx]
    full_text = para.text
    if start < 0 or end > len(full_text) or start >= end:
        raise ToolError(f"文本范围 [{start}:{end}] 无效（段落长度 {len(full_text)}）")

    # 简化实现：在指定范围添加格式化的 run
    runs = para.runs
    current_pos = 0
    for run in runs:
        run_len = len(run.text)
        run_start = current_pos
        run_end = current_pos + run_len

        # 检查是否有重叠
        if run_end > start and run_start < end:
            if bold is not None:
                run.font.bold = bold
            if italic is not None:
                run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
            if font:
                run.font.name = font

        current_pos = run_end

    _save_document(doc, filename, session_id)
    return {"filename": filename, "paragraph_idx": paragraph_idx, "range": [start, end]}


def word_format_table(
    filename: str,
    table_idx: int,
    border_style: str = "single",
    header_row: bool = True,
    shading: str | None = None,
    session_id: str | None = None,
) -> dict[str, Any]:
    """格式化表格。"""
    doc = _get_document(filename, session_id)
    if table_idx < 0 or table_idx >= len(doc.tables):
        raise ToolError(f"表格索引超出范围: {table_idx}（共 {len(doc.tables)} 个表格）")

    table = doc.tables[table_idx]

    # 设置表头加粗
    if header_row and len(table.rows) > 0:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    # 设置边框
    if border_style:
        tbl_pr = table._tbl.tblPr
        existing_borders = tbl_pr.find(qn("w:tblBorders"))
        if existing_borders is not None:
            tbl_pr.remove(existing_borders)
        borders = etree.SubElement(tbl_pr, qn("w:tblBorders"))
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border_elem = etree.SubElement(borders, qn(f"w:{edge}"))
            border_elem.set(qn("w:val"), border_style)
            border_elem.set(qn("w:sz"), "4")
            border_elem.set(qn("w:space"), "0")
            border_elem.set(qn("w:color"), "000000")

    # 设置底纹
    if shading:
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = etree.SubElement(tc_pr, qn("w:shd"))
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), shading)

    _save_document(doc, filename, session_id)
    return {
        "filename": filename,
        "table_idx": table_idx,
        "border_style": border_style,
        "shading": shading,
    }


def word_search_replace(
    filename: str,
    find_text: str,
    replace_text: str,
    session_id: str | None = None,
) -> dict[str, Any]:
    """搜索替换。"""
    InputValidator.validate_text_length(find_text)
    InputValidator.validate_text_length(replace_text)
    doc = _get_document(filename, session_id)
    replaced_count = 0

    # 段落中替换
    for para in doc.paragraphs:
        for run in para.runs:
            if find_text in run.text:
                count = run.text.count(find_text)
                run.text = run.text.replace(find_text, replace_text)
                replaced_count += count

    # 表格中替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if find_text in run.text:
                            count = run.text.count(find_text)
                            run.text = run.text.replace(find_text, replace_text)
                            replaced_count += count

    _save_document(doc, filename, session_id)
    return {"filename": filename, "find": find_text, "replaced_count": replaced_count}


def word_delete_paragraph(
    filename: str,
    paragraph_idx: int,
    session_id: str | None = None,
) -> dict[str, Any]:
    """删除段落。"""
    doc = _get_document(filename, session_id)
    if paragraph_idx < 0 or paragraph_idx >= len(doc.paragraphs):
        raise ToolError(f"段落索引超出范围: {paragraph_idx}（共 {len(doc.paragraphs)} 段）")

    para = doc.paragraphs[paragraph_idx]
    para_element = para._element
    para_element.getparent().remove(para_element)

    _save_document(doc, filename, session_id)
    return {"filename": filename, "deleted_paragraph_idx": paragraph_idx}


def word_create_style(
    filename: str,
    style_name: str,
    font: str = "Calibri",
    size: int = 11,
    color: str = "000000",
    bold: bool = False,
    session_id: str | None = None,
) -> dict[str, Any]:
    """创建自定义样式。"""
    doc = _get_document(filename, session_id)
    styles = doc.styles

    if style_name in [s.name for s in styles]:
        raise ToolError(f"样式已存在: {style_name}")

    new_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    new_style.font.name = font
    new_style.font.size = Pt(size)
    new_style.font.color.rgb = RGBColor.from_string(color)
    new_style.font.bold = bold

    _save_document(doc, filename, session_id)
    return {"filename": filename, "style_name": style_name, "font": font, "size": size}


# ==================== 5.1.4 分析工具（1 个） ====================


def word_extract_tables(
    filename: str,
    format: str = "json",
    session_id: str | None = None,
) -> dict[str, Any]:
    """提取所有表格数据。"""
    InputValidator.validate_choice(format, ["json", "csv"], "format")
    doc = _get_document(filename, session_id)

    tables_data: list[Any] = []
    for idx, table in enumerate(doc.tables):
        rows_data: list[list[str]] = []
        for row in table.rows:
            rows_data.append([cell.text for cell in row.cells])

        if format == "csv":
            csv_lines: list[str] = []
            for row in rows_data:
                csv_lines.append(",".join(f'"{c}"' for c in row))
            tables_data.append({"table_idx": idx, "csv": "\n".join(csv_lines)})
        else:
            tables_data.append({"table_idx": idx, "data": rows_data})

    return {
        "filename": filename,
        "format": format,
        "tables": tables_data,
        "count": len(tables_data),
    }
