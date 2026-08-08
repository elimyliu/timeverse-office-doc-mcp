"""测试 Word handler 工具。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from timeverse_office_doc_mcp.common.error_handler import ToolError
from timeverse_office_doc_mcp.handlers import word_handler


@pytest.fixture
def workspace(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    """创建临时工作空间并更新 PathGuard 白名单。"""
    ws = tmp_path / "workspace"
    ws.mkdir()
    # 更新全局 path_guard 的白名单
    from timeverse_office_doc_mcp.common.path_guard import path_guard

    monkeypatch.setattr(path_guard, "allowed_dirs", [tmp_path.resolve()])
    return ws


@pytest.fixture
def doc_path(workspace: Path) -> str:
    """返回测试文档路径。"""
    return str(workspace / "test.docx")


class TestWordCreateDocument:
    """测试文档创建。"""

    def test_create_empty(self, doc_path: str) -> None:
        """创建空文档。"""
        result = word_handler.word_create_document(doc_path, title="测试", author="pytest")
        assert result["filename"] == str(Path(doc_path).resolve())
        assert result["title"] == "测试"
        assert result["author"] == "pytest"
        assert Path(doc_path).exists()

    def test_create_and_read_back(self, doc_path: str) -> None:
        """创建后能被 python-docx 读回。"""
        word_handler.word_create_document(doc_path, title="回读测试")
        doc = Document(doc_path)
        assert doc.core_properties.title == "回读测试"


class TestWordAddHeading:
    """测试添加标题。"""

    def test_add_heading(self, doc_path: str) -> None:
        """添加标题。"""
        word_handler.word_create_document(doc_path)
        result = word_handler.word_add_heading(doc_path, text="第一章", level=1)
        assert result["text"] == "第一章"
        assert result["level"] == 1
        doc = Document(doc_path)
        assert doc.paragraphs[0].text == "第一章"

    def test_invalid_level(self, doc_path: str) -> None:
        """无效标题级别应报错。"""
        word_handler.word_create_document(doc_path)
        with pytest.raises(ToolError, match="标题级别"):
            word_handler.word_add_heading(doc_path, text="test", level=99)

    def test_add_heading_with_format(self, doc_path: str) -> None:
        """标题支持对齐、字号与段前/段后间距。"""
        word_handler.word_create_document(doc_path)
        result = word_handler.word_add_heading(
            doc_path,
            text="封面标题",
            level=0,
            align="center",
            font_size=26,
            space_before=100,
            space_after=50,
        )
        assert result["align"] == "center"
        assert result["font_size"] == 26
        doc = Document(doc_path)
        para = doc.paragraphs[0]
        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert para.runs[0].font.size is not None
        assert para.paragraph_format.space_before is not None
        assert para.paragraph_format.space_after is not None


class TestWordAddParagraph:
    """测试添加段落。"""

    def test_add_paragraph(self, doc_path: str) -> None:
        """添加段落。"""
        word_handler.word_create_document(doc_path)
        result = word_handler.word_add_paragraph(doc_path, text="段落内容", bold=True, font_size=14)
        assert result["text"] == "段落内容"
        doc = Document(doc_path)
        para = doc.paragraphs[-1]
        assert para.text == "段落内容"
        assert para.runs[0].font.bold is True
        assert para.runs[0].font.size is not None

    def test_page_break(self, doc_path: str) -> None:
        """插入分页符。"""
        word_handler.word_create_document(doc_path)
        result = word_handler.word_add_paragraph(doc_path, page_break=True)
        assert result["filename"] == doc_path or Path(doc_path).name in result["filename"]

    def test_alignment_and_spacing(self, doc_path: str) -> None:
        """段落支持对齐与段前/段后间距（封面排版）。"""
        word_handler.word_create_document(doc_path)
        word_handler.word_add_paragraph(
            doc_path, text="居中段落", align="center", space_before=10, space_after=20
        )
        doc = Document(doc_path)
        para = doc.paragraphs[-1]
        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert para.paragraph_format.space_before is not None
        assert para.paragraph_format.space_after is not None

    def test_invalid_align(self, doc_path: str) -> None:
        """无效对齐方式应报错。"""
        word_handler.word_create_document(doc_path)
        with pytest.raises(ToolError, match="对齐方式"):
            word_handler.word_add_paragraph(doc_path, text="测试", align="middle")

    def test_invalid_spacing(self, doc_path: str) -> None:
        """负间距应报错。"""
        word_handler.word_create_document(doc_path)
        with pytest.raises(ToolError, match="space_before"):
            word_handler.word_add_paragraph(doc_path, text="测试", space_before=-5)


class TestWordAddCover:
    """测试添加封面。"""

    def test_add_cover(self, doc_path: str) -> None:
        """添加居中版式封面并自动分页。"""
        word_handler.word_create_document(doc_path)
        result = word_handler.word_add_cover(
            doc_path,
            title="2026年度项目总结",
            subtitle="技术部年度汇报",
            author="张三",
            date="2026-08-08",
            org="某科技有限公司",
        )
        assert result["title"] == "2026年度项目总结"
        assert result["page_break_added"] is True
        doc = Document(doc_path)
        # 标题段（跳过顶部 2 个空行）：居中、26pt、加粗
        title_para = doc.paragraphs[2]
        assert title_para.text.strip() == "2026年度项目总结"
        assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert title_para.runs[0].font.size == Pt(26)
        assert title_para.runs[0].font.bold is True
        # 副标题段存在
        assert doc.paragraphs[3].text.strip() == "技术部年度汇报"
        # 底部信息存在
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "张三" in texts and "某科技有限公司" in texts
        # 封面后存在分页符
        page_breaks = doc.element.body.findall(".//" + qn("w:br"))
        assert any(br.get(qn("w:type")) == "page" for br in page_breaks)

    def test_empty_title(self, doc_path: str) -> None:
        """空标题应报错。"""
        word_handler.word_create_document(doc_path)
        with pytest.raises(ToolError, match="title"):
            word_handler.word_add_cover(doc_path, title="   ")


class TestWordAddTable:
    """测试添加表格。"""

    def test_add_table_with_data(self, doc_path: str) -> None:
        """添加带数据的表格。"""
        word_handler.word_create_document(doc_path)
        data = [["姓名", "年龄"], ["张三", "25"], ["李四", "30"]]
        result = word_handler.word_add_table(doc_path, rows=3, cols=2, data=data, has_header=True)
        assert result["rows"] == 3
        assert result["cols"] == 2
        doc = Document(doc_path)
        table = doc.tables[0]
        assert table.cell(0, 0).text == "姓名"
        assert table.cell(1, 1).text == "25"

    def test_add_empty_table(self, doc_path: str) -> None:
        """添加空表格。"""
        word_handler.word_create_document(doc_path)
        result = word_handler.word_add_table(doc_path, rows=2, cols=3)
        assert result["rows"] == 2
        assert result["cols"] == 3


class TestWordSearchReplace:
    """测试搜索替换。"""

    def test_basic_replace(self, doc_path: str) -> None:
        """基本搜索替换。"""
        word_handler.word_create_document(doc_path)
        word_handler.word_add_paragraph(doc_path, text="hello world hello")
        result = word_handler.word_search_replace(doc_path, find_text="hello", replace_text="你好")
        assert result["replaced_count"] == 2
        doc = Document(doc_path)
        assert "你好" in doc.paragraphs[-1].text
        assert "hello" not in doc.paragraphs[-1].text


class TestWordGetInfoDetailed:
    """测试文档信息（含结构分析）。"""

    def test_basic_info(self, doc_path: str) -> None:
        """获取基本元信息。"""
        word_handler.word_create_document(doc_path, title="测试", author="pytest")
        result = word_handler.word_get_info(doc_path)
        assert result["title"] == "测试"
        assert result["paragraph_count"] >= 0

    def test_detailed(self, doc_path: str) -> None:
        """获取详细结构分析。"""
        word_handler.word_create_document(doc_path)
        word_handler.word_add_heading(doc_path, text="标题1", level=1)
        word_handler.word_add_paragraph(doc_path, text="段落内容")
        word_handler.word_add_table(doc_path, rows=2, cols=2, data=[["A", "B"], ["1", "2"]])

        result = word_handler.word_get_info(doc_path, detailed=True)
        assert result["paragraph_count"] >= 2
        assert result["table_count"] == 1
        assert result["heading_levels"].get(1) == 1
        assert result["table_details"][0]["rows"] == 2


class TestWordExtractText:
    """测试文本提取。"""

    def test_extract_text(self, doc_path: str) -> None:
        """提取全文文本。"""
        word_handler.word_create_document(doc_path)
        word_handler.word_add_paragraph(doc_path, text="Hello World")
        result = word_handler.word_extract_text(doc_path, extract_type="text")
        assert "Hello World" in result["text"]

    def test_extract_outline(self, doc_path: str) -> None:
        """提取大纲。"""
        word_handler.word_create_document(doc_path)
        word_handler.word_add_heading(doc_path, text="第一章", level=1)
        word_handler.word_add_heading(doc_path, text="第二章", level=2)
        result = word_handler.word_extract_text(doc_path, extract_type="outline")
        assert len(result["outline"]) == 2
        assert result["outline"][0]["text"] == "第一章"

    def test_extract_all(self, doc_path: str) -> None:
        """提取全文和大纲。"""
        word_handler.word_create_document(doc_path)
        word_handler.word_add_heading(doc_path, text="标题", level=1)
        word_handler.word_add_paragraph(doc_path, text="内容")
        result = word_handler.word_extract_text(doc_path, extract_type="all")
        assert "text" in result
        assert "outline" in result


class TestWordExtractTables:
    """测试表格提取。"""

    def test_extract_json(self, doc_path: str) -> None:
        """提取表格数据（JSON）。"""
        word_handler.word_create_document(doc_path)
        data = [["姓名", "年龄"], ["张三", "25"]]
        word_handler.word_add_table(doc_path, rows=2, cols=2, data=data)
        result = word_handler.word_extract_tables(doc_path, format="json")
        assert result["count"] == 1
        assert result["tables"][0]["data"] == data

    def test_extract_csv(self, doc_path: str) -> None:
        """提取表格数据（CSV）。"""
        word_handler.word_create_document(doc_path)
        data = [["姓名", "年龄"], ["张三", "25"]]
        word_handler.word_add_table(doc_path, rows=2, cols=2, data=data)
        result = word_handler.word_extract_tables(doc_path, format="csv")
        assert result["count"] == 1
        assert "姓名" in result["tables"][0]["csv"]


class TestWordSession:
    """测试 Session 内存编辑模式。"""

    def test_session_workflow(self, doc_path: str) -> None:
        """Session 模式工作流：打开 -> 编辑多次 -> 保存。"""
        from timeverse_office_doc_mcp.common.session import session_manager

        # 创建文档并打开 Session
        word_handler.word_create_document(doc_path)
        doc = Document(doc_path)
        sid = session_manager.open_session(doc_path, "word", doc)

        # 在 Session 中多次编辑（不写磁盘）
        word_handler.word_add_heading(doc_path, text="标题", level=1, session_id=sid)
        word_handler.word_add_paragraph(doc_path, text="段落1", session_id=sid)
        word_handler.word_add_paragraph(doc_path, text="段落2", session_id=sid)

        # 确认 Session 状态
        session = session_manager.get_session(sid)
        assert session.modified is True

        # 保存到磁盘
        session_manager.save_session(sid)
        doc.save(doc_path)  # 实际写盘
        session_manager.close_session(sid)

        # 读回验证
        doc2 = Document(doc_path)
        texts = [p.text for p in doc2.paragraphs]
        assert "标题" in texts
        assert "段落1" in texts
        assert "段落2" in texts
